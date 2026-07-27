from __future__ import annotations

from io import BytesIO
import json
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from backend.app.services.document_exports import (
    expand_embedded_result_json,
    markdown_to_docx,
    output_to_markdown,
)


SAMPLE_MARKDOWN = """# 死锁的必要条件

**结论：四个条件必须同时成立。**

## 四个条件

1. 互斥条件
2. 请求与保持
3. 不可剥夺
4. 循环等待

> 破坏其中任意一个条件即可预防死锁。

| 条件 | 含义 | 建议 |
|---|---|---|
| 互斥 | 资源不能共享 | 谨慎评估 |
| 循环等待 | 形成等待闭环 | 固定加锁顺序 |

```python
with first_lock:
    acquire(second_lock)
```
"""


def test_output_to_markdown_unwraps_workflow_result_objects():
    assert output_to_markdown('{"result":"## 可读结果\\n\\n正文"}') == "## 可读结果\n\n正文"
    assert output_to_markdown({"answer": "**完成**"}) == "**完成**"
    assert output_to_markdown(
        json.dumps({"output": {"result": "# 最终稿\n\n正文"}}, ensure_ascii=False)
    ) == "# 最终稿\n\n正文"


def test_output_to_markdown_does_not_export_a_full_node_context_as_json():
    snapshot = {
        "input": {"task": "写综述"},
        "researcher": {"output": "检索中间结果"},
        "artifact": {"output": "# 终稿\n\n可读内容"},
        "output": {"result": "# 终稿\n\n可读内容"},
    }
    result = output_to_markdown(json.dumps(snapshot, ensure_ascii=False))
    assert result == "# 终稿\n\n可读内容"
    assert "```json" not in result


def test_historical_artifact_json_envelope_is_expanded_before_word_export():
    historical = """# 工作流 · 第 1 轮产出

## 执行结果

```json
{"result":"# Mesh Quality Review\\n\\n## Abstract\\n\\nReadable body.\\n\\n| Metric | Value |\\n|---|---|\\n| Jacobian | 0.8 |"}
```
"""
    expanded = expand_embedded_result_json(historical)
    assert "```json" not in expanded
    assert "\\n" not in expanded
    assert "## Abstract" in expanded

    document = Document(BytesIO(markdown_to_docx("Mesh Review", historical)))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert '{"result"' not in text
    assert "Readable body." in text
    assert any(paragraph.text == "Abstract" for paragraph in document.paragraphs)
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "Jacobian"


def test_markdown_to_docx_preserves_document_structure_and_word_geometry():
    payload = markdown_to_docx("死锁分析", SAMPLE_MARKDOWN)
    assert payload.startswith(b"PK")

    document = Document(BytesIO(payload))
    assert document.core_properties.title == "死锁分析"
    assert document.paragraphs[0].style.name == "Title"
    assert document.paragraphs[0].text == "死锁的必要条件"
    assert any(paragraph.style.name == "Heading 1" for paragraph in document.paragraphs)
    assert any("破坏其中任意一个条件" in paragraph.text for paragraph in document.paragraphs)
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 3
    assert [cell.text for cell in document.tables[0].rows[0].cells] == ["条件", "含义", "建议"]

    section = document.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.left_margin.inches, 2) == 1.0
    assert round(section.right_margin.inches, 2) == 1.0

    with ZipFile(BytesIO(payload)) as archive:
        numbering = archive.read("word/numbering.xml").decode("utf-8")
    assert 'w:numFmt w:val="decimal"' in numbering
    assert 'w:numFmt w:val="bullet"' in numbering
    table_properties = document.tables[0]._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    assert table_width.get(qn("w:w")) == "9360"
    assert table_width.get(qn("w:type")) == "dxa"
    assert table_indent.get(qn("w:w")) == "120"
    assert table_indent.get(qn("w:type")) == "dxa"
