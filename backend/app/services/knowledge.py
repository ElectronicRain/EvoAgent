from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader
from sqlalchemy import func, select, text
from sqlalchemy.ext.asyncio import AsyncSession

from ..models import KnowledgeBase, KnowledgeChunk, KnowledgeDocument
from .common import audit


def chunk_text(content: str, chunk_size: int = 900, overlap: int = 120) -> list[str]:
    clean = re.sub(r"\r\n?", "\n", content).strip()
    if not clean:
        return []
    paragraphs = [item.strip() for item in re.split(r"\n{2,}", clean) if item.strip()]
    chunks: list[str] = []
    buffer = ""
    for paragraph in paragraphs:
        if len(buffer) + len(paragraph) + 2 <= chunk_size:
            buffer = f"{buffer}\n\n{paragraph}".strip()
            continue
        if buffer:
            chunks.append(buffer)
        if len(paragraph) <= chunk_size:
            buffer = paragraph
        else:
            start = 0
            while start < len(paragraph):
                chunks.append(paragraph[start : start + chunk_size])
                start += max(1, chunk_size - overlap)
            buffer = ""
    if buffer:
        chunks.append(buffer)
    return chunks


def extract_document(filename: str, data: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages), "application/pdf"
    if suffix == ".docx":
        document = Document(io.BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if suffix in {".txt", ".md", ".csv"}:
        return data.decode("utf-8", errors="replace"), "text/plain"
    raise ValueError("当前支持 PDF、DOCX、TXT、MD 和 CSV")


class KnowledgeService:
    async def add_document(
        self,
        db: AsyncSession,
        knowledge_base_id: str,
        *,
        title: str,
        content: str,
        source: str,
        mime_type: str = "text/plain",
    ) -> KnowledgeDocument:
        knowledge_base = await db.get(KnowledgeBase, knowledge_base_id)
        if not knowledge_base:
            raise LookupError("知识库不存在")
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
        duplicate = await db.scalar(
            select(KnowledgeDocument).where(
                KnowledgeDocument.knowledge_base_id == knowledge_base_id,
                KnowledgeDocument.content_hash == content_hash,
            )
        )
        if duplicate:
            return duplicate
        document = KnowledgeDocument(
            knowledge_base_id=knowledge_base_id,
            title=title,
            source=source,
            mime_type=mime_type,
            content_hash=content_hash,
            char_count=len(content),
        )
        db.add(document)
        await db.flush()
        chunks = chunk_text(content)
        for index, value in enumerate(chunks):
            chunk = KnowledgeChunk(
                document_id=document.id,
                knowledge_base_id=knowledge_base_id,
                chunk_index=index,
                title=title,
                content=value,
                citation=f"{title}，片段 {index + 1}，来源：{source}",
            )
            db.add(chunk)
            await db.flush()
            await db.execute(
                text(
                    "INSERT INTO knowledge_chunks_fts(chunk_id, title, content) "
                    "VALUES (:id, :title, :content)"
                ),
                {"id": chunk.id, "title": title, "content": value},
            )
        knowledge_base.document_count += 1
        await audit(
            db,
            "knowledge.document_added",
            "knowledge_document",
            document.id,
            {"title": title, "chunks": len(chunks)},
        )
        return document

    async def search(
        self,
        db: AsyncSession,
        query: str,
        knowledge_base_ids: list[str] | None = None,
        top_k: int = 5,
    ) -> list[dict[str, Any]]:
        ids = knowledge_base_ids or []
        rows = []
        try:
            sql = (
                "SELECT c.id, c.title, c.content, c.citation, c.knowledge_base_id, "
                "bm25(knowledge_chunks_fts) AS rank "
                "FROM knowledge_chunks_fts f JOIN knowledge_chunks c ON c.id=f.chunk_id "
                "WHERE knowledge_chunks_fts MATCH :query"
            )
            parameters: dict[str, Any] = {"query": query, "limit": top_k}
            if ids:
                placeholders = ",".join(f":kb{i}" for i in range(len(ids)))
                sql += f" AND c.knowledge_base_id IN ({placeholders})"
                parameters.update({f"kb{i}": value for i, value in enumerate(ids)})
            sql += " ORDER BY rank LIMIT :limit"
            rows = (await db.execute(text(sql), parameters)).mappings().all()
        except Exception:
            rows = []
        if not rows:
            statement = select(KnowledgeChunk)
            if ids:
                statement = statement.where(KnowledgeChunk.knowledge_base_id.in_(ids))
            chunks = (await db.scalars(statement.limit(2000))).all()
            query_chars = set(re.sub(r"\s+", "", query))

            def score(item: KnowledgeChunk) -> float:
                content = re.sub(r"\s+", "", item.content)
                overlap = len(query_chars & set(content)) / max(1, len(query_chars))
                phrase_bonus = 1.0 if query in item.content else 0.0
                return phrase_bonus + overlap

            ranked = sorted(chunks, key=score, reverse=True)
            chunks = [item for item in ranked[:top_k] if score(item) > 0]
            rows = [
                {
                    "id": item.id,
                    "title": item.title,
                    "content": item.content,
                    "citation": item.citation,
                    "knowledge_base_id": item.knowledge_base_id,
                    "rank": -score(item),
                }
                for item in chunks
            ]
        return [dict(item) for item in rows]

    async def stats(self, db: AsyncSession) -> dict[str, int]:
        return {
            "bases": int(await db.scalar(select(func.count(KnowledgeBase.id))) or 0),
            "documents": int(await db.scalar(select(func.count(KnowledgeDocument.id))) or 0),
            "chunks": int(await db.scalar(select(func.count(KnowledgeChunk.id))) or 0),
        }


knowledge_service = KnowledgeService()
