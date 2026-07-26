from __future__ import annotations

from html import unescape
from io import BytesIO
import json
import re
from typing import Any
from urllib.parse import quote

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


_INLINE_PATTERN = re.compile(
    r"(!?\[[^\]]+\]\([^)]+\)|\*\*[^*\n]+\*\*|__[^_\n]+__|"
    r"~~[^~\n]+~~|`[^`\n]+`|\*[^*\n]+\*|_[^_\n]+_)"
)
_TABLE_SEPARATOR = re.compile(r"^:?-{3,}:?$")
_INVALID_FILENAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_JSON_FENCE = re.compile(r"```json\s*\n([\s\S]*?)\n```", re.IGNORECASE)

BODY_COLOR = RGBColor(38, 55, 70)
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(103, 121, 136)
TABLE_HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "172B3A"
CODE_COLOR = RGBColor(235, 242, 247)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def safe_docx_filename(title: str) -> str:
    stem = _INVALID_FILENAME.sub("-", str(title or "工作流成果")).strip(" .-")
    stem = re.sub(r"\s+", " ", stem)[:90] or "工作流成果"
    return f"{stem}.docx"


def content_disposition(filename: str) -> str:
    return f"attachment; filename=workflow-result.docx; filename*=UTF-8''{quote(filename)}"


def output_to_markdown(value: str | dict[str, Any] | list[Any] | Any) -> str:
    current = value
    for _ in range(3):
        if not isinstance(current, str):
            break
        source = current.strip()
        if not source or source[:1] not in "[{\"":
            return current
        try:
            current = json.loads(source)
        except (TypeError, json.JSONDecodeError):
            return current
    if isinstance(current, str):
        return current
    if isinstance(current, dict):
        preferred = ("result", "output", "content", "answer", "markdown", "text", "document")
        for key in preferred:
            candidate = current.get(key)
            if isinstance(candidate, str) and candidate.strip():
                return candidate
        sections = []
        for key, candidate in current.items():
            if isinstance(candidate, str) and candidate.strip():
                sections.append(f"## {key}\n\n{candidate}")
        if sections:
            return "\n\n".join(sections)
    return f"```json\n{json.dumps(current, ensure_ascii=False, indent=2, default=str)}\n```"


def expand_embedded_result_json(markdown: str) -> str:
    """Expand JSON result envelopes stored by earlier workflow versions.

    Historical artifacts wrapped a Markdown answer in a fenced JSON object,
    for example `````json {"result": "# Title\\n..."} `````.  Treating that
    fence as source code is technically valid Markdown, but produces an
    unreadable Word document with visible ``\\n`` escapes.  Only envelopes
    containing a well-known textual result key are expanded; ordinary JSON
    examples remain code blocks.
    """

    def replace(match: re.Match[str]) -> str:
        try:
            payload = json.loads(match.group(1).strip())
        except (TypeError, json.JSONDecodeError):
            return match.group(0)
        if not isinstance(payload, dict):
            return match.group(0)
        for key in ("result", "output", "content", "answer", "markdown", "text", "document"):
            value = payload.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
        return match.group(0)

    expanded = str(markdown or "")
    for _ in range(3):
        updated = _JSON_FENCE.sub(replace, expanded)
        if updated == expanded:
            break
        expanded = updated
    return expanded


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style: Any, size: float, color: RGBColor, *, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _configure_document(document: Document, running_title: str) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, 11, BODY_COLOR)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    _set_style_font(title, 24, RGBColor(23, 63, 96), bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 18, 10),
        "Heading 2": (13, HEADING_BLUE, 14, 7),
        "Heading 3": (12, HEADING_DARK_BLUE, 10, 5),
        "Heading 4": (11, HEADING_DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        _set_style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run(f"EvoAgent  ·  {running_title[:72]}")
    _set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    prefix = footer.add_run("EvoAgent 工作流成果  ·  第 ")
    _set_run_font(prefix, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instruction, separate, text, end])
    footer._p.append(field_run)
    suffix = footer.add_run(" 页")
    _set_run_font(suffix, size=8, color=MUTED)


def _add_numbering_definition(document: Document, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId"), "0"))
        for item in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), f"%{level + 1}." if ordered else "•")
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(540 + level * 360))
        indent.set(qn("w:hanging"), "280")
        p_pr.extend([tabs, indent])
        lvl.extend([start, fmt, text, justification, p_pr])
        abstract.append(lvl)
    numbering.append(abstract)

    num_ids = [int(item.get(qn("w:numId"), "0")) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph: Any, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(max(level, 0), 2)))
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, number])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def _add_hyperlink(paragraph: Any, label: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1769C2")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph: Any, value: str) -> None:
    source = unescape(re.sub(r"<br\s*/?>", "\n", str(value), flags=re.IGNORECASE))
    for line_index, line in enumerate(source.split("\n")):
        if line_index:
            paragraph.add_run().add_break()
        position = 0
        for match in _INLINE_PATTERN.finditer(line):
            if match.start() > position:
                _set_run_font(paragraph.add_run(line[position : match.start()]), size=11, color=BODY_COLOR)
            token = match.group(0)
            link_match = re.fullmatch(r"(!?)\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                prefix, label, url = link_match.groups()
                _add_hyperlink(paragraph, f"图片：{label}" if prefix else label, url)
            elif token.startswith(("**", "__")):
                _set_run_font(paragraph.add_run(token[2:-2]), size=11, color=BODY_COLOR, bold=True)
            elif token.startswith("~~"):
                run = paragraph.add_run(token[2:-2])
                _set_run_font(run, size=11, color=BODY_COLOR)
                run.font.strike = True
            elif token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                _set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.5, color=RGBColor(156, 48, 81))
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "EDF2F6")
                run._element.get_or_add_rPr().append(shading)
            else:
                _set_run_font(paragraph.add_run(token[1:-1]), size=11, color=BODY_COLOR, italic=True)
            position = match.end()
        if position < len(line):
            _set_run_font(paragraph.add_run(line[position:]), size=11, color=BODY_COLOR)


def _split_table_row(line: str) -> list[str]:
    source = line.strip().strip("|")
    return [cell.strip().replace(r"\|", "|") for cell in re.split(r"(?<!\\)\|", source)]


def _is_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines) or "|" not in lines[index]:
        return False
    separators = _split_table_row(lines[index + 1])
    return bool(separators) and all(_TABLE_SEPARATOR.fullmatch(item.replace(" ", "")) for item in separators)


def _table_widths(rows: list[list[str]], columns: int) -> list[int]:
    weights = []
    for index in range(columns):
        longest = max((len(row[index]) if index < len(row) else 0 for row in rows), default=1)
        weights.append(min(36, max(8, longest)))
    total = sum(weights) or columns
    raw = [max(720, round(CONTENT_WIDTH_DXA * value / total)) for value in weights]
    scale = CONTENT_WIDTH_DXA / sum(raw)
    widths = [round(value * scale) for value in raw]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(document: Document, header: list[str], body: list[list[str]]) -> None:
    columns = max(1, len(header))
    rows = [header] + [row + [""] * (columns - len(row)) for row in body]
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            _add_inline(paragraph, values[column_index] if column_index < len(values) else "")
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), TABLE_HEADER_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = HEADING_DARK_BLUE
        if row_index == 0:
            header_property = OxmlElement("w:tblHeader")
            header_property.set(qn("w:val"), "true")
            table.rows[0]._tr.get_or_add_trPr().append(header_property)
    _set_table_geometry(table, _table_widths(rows, columns))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBDDE9")
    borders.append(bottom)
    p_pr.append(borders)


def _add_code_block(document: Document, code: str, language: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.1
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    paragraph._p.get_or_add_pPr().append(shading)
    if language:
        language_run = paragraph.add_run(f"{language}\n")
        _set_run_font(language_run, name="Consolas", size=8, color=RGBColor(133, 190, 222), bold=True)
    run = paragraph.add_run(code.rstrip())
    _set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9, color=CODE_COLOR)


def _add_blockquote(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.extend([shading, borders])
    _add_inline(paragraph, "\n".join(lines))
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(73, 105, 126)


def markdown_to_docx(title: str, markdown: str) -> bytes:
    document = Document()
    _configure_document(document, title)
    bullet_num_id = _add_numbering_definition(document, ordered=False)
    ordered_num_id = _add_numbering_definition(document, ordered=True)

    source = expand_embedded_result_json(str(markdown or ""))
    source = source.replace("\r\n", "\n").replace("\r", "\n").strip()
    fence_match = re.fullmatch(r"```(?:markdown|md)?\s*\n([\s\S]*?)\n```", source, re.IGNORECASE)
    if fence_match:
        source = fence_match.group(1).strip()
    lines = source.split("\n") if source else []
    has_markdown_title = any(re.match(r"^#\s+\S", line.strip()) for line in lines[:5])
    if not has_markdown_title:
        paragraph = document.add_paragraph(style="Title")
        _add_inline(paragraph, title or "工作流成果")

    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        _add_inline(paragraph, "\n".join(item.strip() for item in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            _add_code_block(document, "\n".join(code_lines), language)
            index += 1
            continue
        if _is_table(lines, index):
            flush_paragraph()
            header = _split_table_row(lines[index])
            index += 2
            body = []
            while index < len(lines) and lines[index].strip() and "|" in lines[index]:
                body.append(_split_table_row(lines[index]))
                index += 1
            _add_table(document, header, body)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            style = "Title" if level == 1 else f"Heading {min(level - 1, 4)}"
            paragraph = document.add_paragraph(style=style)
            _add_inline(paragraph, heading.group(2))
            index += 1
            continue
        if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,})", stripped):
            flush_paragraph()
            _add_rule(document)
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[index]))
                index += 1
            _add_blockquote(document, quote_lines)
            continue
        list_match = re.match(r"^(\s*)([-+*]|\d+[.)])\s+(.+)$", line)
        if list_match:
            flush_paragraph()
            level = min(len(list_match.group(1).replace("\t", "    ")) // 2, 2)
            ordered = bool(re.match(r"\d", list_match.group(2)))
            paragraph = document.add_paragraph()
            _apply_numbering(paragraph, ordered_num_id if ordered else bullet_num_id, level)
            _add_inline(paragraph, list_match.group(3))
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()

    properties = document.core_properties
    properties.title = title
    properties.subject = "EvoAgent 工作流产出"
    properties.author = "EvoAgent"
    properties.keywords = "EvoAgent, Workflow, Markdown"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
